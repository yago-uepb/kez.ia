import io

import docx
import pymupdf


class ContentExtractor:
    def __init__(self):
        self.EXTRACTORS = {
            "application/pdf": self.extract_pdf_text,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self.extract_docx_text,
        }


    def extract_pdf_text(self, content):
        with pymupdf.open(stream=content, filetype="pdf") as document:
            parts = []
            for page in document:
                parts.append(page.get_text())
                for table in page.find_tables():
                    parts.append(table.to_markdown())  # preserva estrutura de linhas/colunas

        return "\n".join(parts).strip()


    def _extract_docx_tables(self, document):
        parts = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))

        return "\n".join(parts)


    def extract_docx_text(self, content):
        document = docx.Document(io.BytesIO(content))
        paragraphs = [p.text for p in document.paragraphs]

        tables_text = self._extract_docx_tables(document)

        return "\n".join(paragraphs + [tables_text]).strip()